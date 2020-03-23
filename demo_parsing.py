from docx import Document
import json
import sys

def load(string):
    """Loads string into dictionary, resiliently to smart quotes."""
    return json.loads(string.replace(u"\u201c", "\"").replace(u"\u201d", "\""))


def single_parse(doc, style_name="requirement"):
    """Parse docx document paragraphs with style_name into dicts."""
    dicts = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name == style_name:
            try:
                dicts.append(load(paragraph.text))
            except:
                print("There was a JSON parsing error while parsing: ")
                print(paragraph)


def multiple_parse(doc, style_name="requirement"):
    """Parse docx document paragraphs.

    Supports style_name paragraphs with multiple lines."""
    meta_paragraphs = [""]
    dicts = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name == style_name:
            meta_paragraphs[-1] += paragraph.text
        else:
            meta_paragraphs.append("")
    for paragraph in meta_paragraphs:
        if paragraph != "":
            try:
                dicts.append(load(paragraph))
            except:
                print("There was a JSON parsing error while parsing: ")
                print(paragraph)
    return dicts


def show_dict(d):
    """Nested view for dictionaries. Sorts keys."""
    print(json.dumps(d, sort_keys=True, indent=4))

if __name__ == "__main__":
    file_name = sys.argv[-1] # last command line argument. All others are ignored.
    doc = Document(file_name)
    dicts= multiple_parse(doc, "Quotations")
    for d in dicts:
        show_dict(d)
